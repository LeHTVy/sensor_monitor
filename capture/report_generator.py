#!/usr/bin/env python3
"""
Report Generator - Creates professional reconnaissance reports
Generates DOCX and PDF reports from reconnaissance scan results
"""

import os
from datetime import datetime
from typing import Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Report storage directory
REPORTS_DIR = os.getenv('REPORTS_DIR', '/app/recon_reports')
os.makedirs(REPORTS_DIR, exist_ok=True)


class ReconReportGenerator:
    """Generates professional reconnaissance reports in DOCX and PDF formats"""
    
    def __init__(self, recon_data: Dict):
        self.recon_data = recon_data
        self.recon_id = recon_data.get('recon_id', 'unknown')
        self.target_ip = recon_data.get('target_ip', 'unknown')
    
    def generate_docx(self) -> str:
        """Generate DOCX report"""
        logger.info(f"[Report] Generating DOCX for recon {self.recon_id}")
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Black Box Reconnaissance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with target
        subtitle = doc.add_heading(f'Target: {self.target_ip}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"Reconnaissance ID: {self.recon_id}")
        doc.add_paragraph(f"Target IP: {self.target_ip}")
        doc.add_paragraph(f"Scan Start: {self.recon_data.get('start_time', 'N/A')}")
        doc.add_paragraph(f"Scan End: {self.recon_data.get('end_time', 'N/A')}")
        doc.add_paragraph(f"Status: {self.recon_data.get('status', 'unknown')}")
        
        tools_used = self.recon_data.get('scan_types', [])
        doc.add_paragraph(f"Tools Used: {', '.join(tools_used)}")
        
        doc.add_page_break()
        
        # Nmap Results
        if 'nmap' in self.recon_data.get('tools', {}):
            doc.add_heading('Nmap Scan Results', level=1)
            self._add_nmap_section(doc, self.recon_data['tools']['nmap'])
            doc.add_page_break()
        
        # Subdomain Enumeration
        subdomain_tools = ['amass', 'subfinder', 'bbot']
        if any(tool in self.recon_data.get('tools', {}) for tool in subdomain_tools):
            doc.add_heading('Subdomain Enumeration', level=1)
            
            for tool in subdomain_tools:
                if tool in self.recon_data.get('tools', {}):
                    self._add_subdomain_section(doc, tool, self.recon_data['tools'][tool])
            
            doc.add_page_break()
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        self._add_recommendations(doc)
        
        # Save document
        filename = f"recon_{self.recon_id}_{self.target_ip.replace('.', '_')}.docx"
        filepath = os.path.join(REPORTS_DIR, filename)
        doc.save(filepath)
        
        logger.info(f"[Report] DOCX saved to {filepath}")
        return filepath
    
    def _add_nmap_section(self, doc, nmap_data: Dict):
        """Add Nmap results to DOCX"""
        
        # Host Discovery
        if 'host_discovery' in nmap_data:
            doc.add_heading('Host Discovery', level=2)
            host_disc = nmap_data['host_discovery']
            host_up = host_disc.get('host_up', False)
            doc.add_paragraph(f"Host Status: {'UP' if host_up else 'DOWN'}")
            doc.add_paragraph()
        
        # Port Scan
        if 'port_scan' in nmap_data:
            doc.add_heading('Port Scan', level=2)
            port_scan = nmap_data['port_scan']
            
            if port_scan.get('status') == 'completed':
                open_ports = port_scan.get('open_ports', [])
                doc.add_paragraph(f"Total Open Ports: {len([p for p in open_ports if p.get('state') == 'open'])}")
                
                if openports := [p for p in open_ports if p.get('state') == 'open']:
                    doc.add_heading('Open Ports', level=3)
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Port'
                    hdr_cells[1].text = 'Protocol'
                    hdr_cells[2].text = 'State'
                    hdr_cells[3].text = 'Service'
                    
                    # Data rows
                    for port in open_ports[:50]:  # Limit to 50 ports
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(port.get('port', ''))
                        row_cells[1].text = port.get('protocol', '')
                        row_cells[2].text = port.get('state', '')
                        row_cells[3].text = port.get('service', 'unknown')
                    
                    doc.add_paragraph()
            else:
                doc.add_paragraph(f"Port scan status: {port_scan.get('status', 'unknown')}")
                if 'error' in port_scan:
                    doc.add_paragraph(f"Error: {port_scan['error']}")
        
        # Service Detection
        if 'service_detection' in nmap_data:
            doc.add_heading('Service Detection', level=2)
            services = nmap_data['service_detection'].get('services', [])
            
            if services:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                # Header
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Port'
                hdr_cells[1].text = 'Service'
                hdr_cells[2].text = 'Product'
                hdr_cells[3].text = 'Version'
                hdr_cells[4].text = 'Extra Info'
                
                # Data
                for svc in services[:30]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(svc.get('port', ''))
                    row_cells[1].text = svc.get('name', '')
                    row_cells[2].text = svc.get('product', '')
                    row_cells[3].text = svc.get('version', '')
                    row_cells[4].text = svc.get('extrainfo', '')
                
                doc.add_paragraph()
        
        # OS Fingerprinting
        if 'os_fingerprinting' in nmap_data:
            doc.add_heading('OS Fingerprinting', level=2)
            os_matches = nmap_data['os_fingerprinting'].get('os_matches', [])
            
            if os_matches:
                for i, match in enumerate(os_matches[:5], 1):
                    doc.add_paragraph(f"{i}. {match.get('name', 'Unknown')} (Accuracy: {match.get('accuracy', 'N/A')}%)")
            else:
                doc.add_paragraph("No OS matches found")
            
            doc.add_paragraph()
    
    def _add_subdomain_section(self, doc, tool: str, tool_data: Dict):
        """Add subdomain enumeration results"""
        doc.add_heading(f'{tool.upper()} Results', level=2)
        
        if tool_data.get('status') == 'completed':
            subdomains = tool_data.get('subdomains', [])
            count = tool_data.get('count', len(subdomains))
            
            doc.add_paragraph(f"Total Subdomains Found: {count}")
            
            if subdomains:
                doc.add_paragraph("Subdomains:")
                for subdomain in subdomains[:50]:  # Limit to 50
                    doc.add_paragraph(f"  • {subdomain}", style='List Bullet')
            doc.add_paragraph()
        else:
            doc.add_paragraph(f"Status: {tool_data.get('status', 'unknown')}")
            if 'error' in tool_data:
                doc.add_paragraph(f"Error: {tool_data['error']}")
    
    def _add_recommendations(self, doc):
        """Add security recommendations based on findings"""
        recommendations = []
        
        # Analyze nmap results for recommendations
        if 'nmap' in self.recon_data.get('tools', {}):
            nmap_data = self.recon_data['tools']['nmap']
            
            if 'port_scan' in nmap_data:
                open_ports = nmap_data['port_scan'].get('open_ports', [])
                open_count = len([p for p in open_ports if p.get('state') == 'open'])
                
                if open_count > 20:
                    recommendations.append("High number of open ports detected. Review and close unnecessary services.")
                
                # Check for commonly exploited ports
                dangerous_ports = [21, 23, 25, 135, 139, 445, 3389]
                found_dangerous = [p['port'] for p in open_ports if p.get('port') in dangerous_ports and p.get('state') == 'open']
                
                if found_dangerous:
                    recommendations.append(f"Potentially dangerous ports found: {', '.join(map(str, found_dangerous))}. Consider restricting access.")
        
        # Add general recommendations
        recommendations.extend([
            "Monitor this IP address for further attack attempts.",
            "Review firewall rules to ensure proper access control.",
            "Implement intrusion detection/prevention systems.",
            "Keep all services and software up to date.",
            "Consider implementing rate limiting and geoblocking for suspicious IPs."
        ])
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")
    
    def generate_pdf(self) -> str:
        """Generate PDF report"""
        logger.info(f"[Report] Generating PDF for recon {self.recon_id}")
        
        filename = f"recon_{self.recon_id}_{self.target_ip.replace('.', '_')}.pdf"
        filepath = os.path.join(REPORTS_DIR, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph('Black Box Reconnaissance Report', title_style))
        story.append(Paragraph(f'Target: {self.target_ip}', styles['Heading2']))
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph('Executive Summary', heading_style))
        summary_data = [
            ['Reconnaissance ID:', self.recon_id],
            ['Target IP:', self.target_ip],
            ['Scan Start:', self.recon_data.get('start_time', 'N/A')],
            ['Scan End:', self.recon_data.get('end_time', 'N/A')],
            ['Status:', self.recon_data.get('status', 'unknown')],
            ['Tools Used:', ', '.join(self.recon_data.get('scan_types', []))]
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(summary_table)
        story.append(PageBreak())
        
        # Nmap Results
        if 'nmap' in self.recon_data.get('tools', {}):
            story.append(Paragraph('Nmap Scan Results', heading_style))
            self._add_nmap_pdf(story, styles, self.recon_data['tools']['nmap'])
            story.append(PageBreak())
        
        # Subdomain Results
        subdomain_tools = ['amass', 'subfinder', 'bbot']
        if any(tool in self.recon_data.get('tools', {}) for tool in subdomain_tools):
            story.append(Paragraph('Subdomain Enumeration', heading_style))
            
            for tool in subdomain_tools:
                if tool in self.recon_data.get('tools', {}):
                    story.append(Paragraph(f'{tool.upper()}', styles['Heading3']))
                    tool_data = self.recon_data['tools'][tool]
                    
                    if tool_data.get('status') == 'completed':
                        count = tool_data.get('count', 0)
                        story.append(Paragraph(f'Subdomains Found: {count}', styles['Normal']))
                        story.append(Spacer(1, 10))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"[Report] PDF saved to {filepath}")
        return filepath
    
    def _add_nmap_pdf(self, story, styles, nmap_data: Dict):
        """Add Nmap results to PDF"""
        
        # Port Scan Results
        if 'port_scan' in nmap_data and nmap_data['port_scan'].get('status') == 'completed':
            open_ports = [p for p in nmap_data['port_scan'].get('open_ports', []) if p.get('state') == 'open']
            
            story.append(Paragraph(f'Open Ports: {len(open_ports)}', styles['Heading3']))
            
            if open_ports[:20]:  # Limit to 20 for PDF
                port_data = [['Port', 'Protocol', 'State', 'Service']]
                for port in open_ports[:20]:
                    port_data.append([
                        str(port.get('port', '')),
                        port.get('protocol', ''),
                        port.get('state', ''),
                        port.get('service', 'unknown')
                    ])
                
                port_table = Table(port_data)
                port_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(port_table)
        
        story.append(Spacer(1, 20))


def generate_report(recon_data: Dict, format: str = 'docx') -> str:
    """
    Generate reconnaissance report
    
    Args:
        recon_data: Full reconnaissance results dictionary
        format: 'docx' or 'pdf'
    
    Returns:
        Path to generated report file
    """
    generator = ReconReportGenerator(recon_data)
    
    if format.lower() == 'pdf':
        return generator.generate_pdf()
    else:
        return generator.generate_docx()
